# main.py
import os
import re
import uuid
from datetime import datetime
from typing import List, Optional

from fastapi import FastAPI, HTTPException, BackgroundTasks
from fastapi.responses import FileResponse
from pydantic import BaseModel, Field
from jinja2 import Environment, BaseLoader, select_autoescape

from motor.motor_asyncio import AsyncIOMotorClient
from bson import ObjectId
from bson.errors import InvalidId

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

from weasyprint import HTML

app = FastAPI(title="Video Synopsis Exporter Engine")

# ---------------------------------------------------------------------------
# Config / Mongo connection
#
# This is the actual document store M5 writes to. We read-only from it here
# -- the export engine never mutates a synopsis.
# ---------------------------------------------------------------------------
MONGO_URI = os.getenv("MONGO_URI", "mongodb://localhost:27017")
MONGO_DB_NAME = os.getenv("MONGO_DB_NAME", "video_synopsis_ai")
MONGO_COLLECTION = os.getenv("MONGO_COLLECTION", "synopses")

mongo_client = AsyncIOMotorClient(MONGO_URI)
synopsis_collection = mongo_client[MONGO_DB_NAME][MONGO_COLLECTION]

OUTPUT_DIR = os.path.join(os.getcwd(), "generated_exports")
os.makedirs(OUTPUT_DIR, exist_ok=True)

BRAND_COLOR = "#1a73e8"
BRAND_COLOR_RGB = RGBColor(0x1A, 0x73, 0xE8)
TEXT_COLOR_RGB = RGBColor(0x20, 0x21, 0x24)
MUTED_RGB = RGBColor(0x80, 0x86, 0x8B)


# ---------------------------------------------------------------------------
# 1. Schema — mirrors M5's exact JSON output, field for field.
#    If M5's payload is missing a required field or has the wrong type,
#    FastAPI returns a 422 with the exact field that failed, instead of
#    a half-rendered PDF.
# ---------------------------------------------------------------------------
class VideoMetadata(BaseModel):
    title: str
    video_url: str
    channel_name: Optional[str] = None
    thumbnail_url: Optional[str] = None


class BasicSummary(BaseModel):
    overall_synopsis: str


class TopicsCovered(BaseModel):
    title: str
    topics: List[str] = Field(default_factory=list)


class TopicBreakdownItem(BaseModel):
    topic: str
    explanation: str


class DetailedSummary(BaseModel):
    key_insights: List[str] = Field(default_factory=list)
    action_items: List[str] = Field(default_factory=list)
    topic_breakdown: List[TopicBreakdownItem] = Field(default_factory=list)


class SummaryBlock(BaseModel):
    basic_summary: BasicSummary
    topics_covered: TopicsCovered
    detailed_summary: DetailedSummary
    closing_note: str


class SynopsisInput(BaseModel):
    video_metadata: VideoMetadata
    summary: SummaryBlock


def _cleanup_file(path: str) -> None:
    try:
        if os.path.exists(path):
            os.remove(path)
    except OSError:
        pass


def _slugify(text: str) -> str:
    slug = re.sub(r"[^a-zA-Z0-9]+", "-", text).strip("-").lower()
    return slug[:60] if slug else "synopsis"


async def _fetch_synopsis(synopsis_id: str) -> SynopsisInput:
    """Pulls the raw document from MongoDB and validates it into our schema."""
    try:
        object_id = ObjectId(synopsis_id)
    except InvalidId:
        raise HTTPException(status_code=400, detail="Invalid synopsis id format.")

    raw_doc = await synopsis_collection.find_one({"_id": object_id})
    if raw_doc is None:
        raise HTTPException(status_code=404, detail="Synopsis not found.")

    raw_doc.pop("_id", None)
    try:
        return SynopsisInput(**raw_doc)
    except Exception as exc:
        # This means the document stored by M5 doesn't match the agreed
        # schema -- surface it clearly instead of failing inside WeasyPrint.
        raise HTTPException(
            status_code=500,
            detail=f"Stored synopsis is malformed and cannot be exported: {exc}",
        )


# ---------------------------------------------------------------------------
# 2. PDF Generation (Jinja2 autoescape + WeasyPrint)
# ---------------------------------------------------------------------------
PDF_TEMPLATE = """
<html>
<head>
<style>
    @page {
        margin: 2.2cm 2cm 2.5cm 2cm;
        @bottom-center {
            content: "Page " counter(page) " of " counter(pages);
            font-size: 9px;
            color: #80868b;
        }
    }
    body { font-family: 'Helvetica Neue', Arial, sans-serif; color: #333333; line-height: 1.6; font-size: 13px; }
    .brand-header { display: flex; justify-content: space-between; align-items: baseline; border-bottom: 3px solid """ + BRAND_COLOR + """; padding-bottom: 6px; margin-bottom: 18px; }
    .brand-header .label { font-size: 11px; letter-spacing: 1px; text-transform: uppercase; color: """ + BRAND_COLOR + """; font-weight: bold; }
    .brand-header .date { font-size: 11px; color: #80868b; }
    h1 { color: #202124; font-size: 22px; margin: 4px 0 14px 0; }
    h2 { color: #202124; margin-top: 24px; font-size: 15px; border-left: 4px solid """ + BRAND_COLOR + """; padding-left: 8px; }
    h3 { color: """ + BRAND_COLOR + """; font-size: 13px; margin: 14px 0 4px 0; }
    .metadata { background-color: #f8f9fa; padding: 12px 14px; border-radius: 6px; margin-bottom: 18px; font-size: 12.5px; border: 1px solid #dadce0; }
    .metadata a { color: """ + BRAND_COLOR + """; text-decoration: none; font-weight: 600; word-break: break-all; }
    .thumb { max-width: 240px; border-radius: 6px; margin-bottom: 14px; display: block; }
    ul { padding-left: 20px; margin-top: 8px; }
    li { margin-bottom: 6px; }
    .topic-card { background: #fafbfc; border: 1px solid #e8eaed; border-radius: 6px; padding: 10px 14px; margin-bottom: 10px; }
    .closing { background: #f0f6ff; border-left: 4px solid """ + BRAND_COLOR + """; padding: 12px 16px; margin-top: 22px; font-style: italic; border-radius: 0 6px 6px 0; }
    .empty { color: #9aa0a6; font-style: italic; font-size: 12px; }
</style>
</head>
<body>
    <div class="brand-header">
        <span class="label">Video Synopsis</span>
        <span class="date">Generated {{ generated_on }}</span>
    </div>

    {% if meta.thumbnail_url %}
    <img class="thumb" src="{{ meta.thumbnail_url }}" />
    {% endif %}

    <h1>{{ meta.title }}</h1>

    <div class="metadata">
        {% if meta.channel_name %}<strong>Channel:</strong> {{ meta.channel_name }}<br/>{% endif %}
        <strong>Source Video:</strong>
        <a href="{{ meta.video_url }}">{{ meta.video_url }}</a>
    </div>

    <h2>Overall Synopsis</h2>
    <p>{{ summary.basic_summary.overall_synopsis }}</p>

    <h2>{{ summary.topics_covered.title }}</h2>
    {% if summary.topics_covered.topics %}
    <ul>{% for t in summary.topics_covered.topics %}<li>{{ t }}</li>{% endfor %}</ul>
    {% else %}<p class="empty">No topics provided.</p>{% endif %}

    <h2>Key Insights</h2>
    {% if summary.detailed_summary.key_insights %}
    <ul>{% for i in summary.detailed_summary.key_insights %}<li>{{ i }}</li>{% endfor %}</ul>
    {% else %}<p class="empty">No insights provided.</p>{% endif %}

    {% if summary.detailed_summary.topic_breakdown %}
    <h2>Topic Breakdown</h2>
    {% for item in summary.detailed_summary.topic_breakdown %}
    <div class="topic-card">
        <h3>{{ item.topic }}</h3>
        <p>{{ item.explanation }}</p>
    </div>
    {% endfor %}
    {% endif %}

    <h2>Action Items</h2>
    {% if summary.detailed_summary.action_items %}
    <ul>{% for a in summary.detailed_summary.action_items %}<li>{{ a }}</li>{% endfor %}</ul>
    {% else %}<p class="empty">No action items provided.</p>{% endif %}

    <div class="closing">{{ summary.closing_note }}</div>
</body>
</html>
"""

_jinja_env = Environment(loader=BaseLoader(), autoescape=select_autoescape(["html", "xml"]))
_pdf_template = _jinja_env.from_string(PDF_TEMPLATE)


def _render_pdf(data: SynopsisInput) -> str:
    html_content = _pdf_template.render(
        meta=data.video_metadata,
        summary=data.summary,
        generated_on=datetime.utcnow().strftime("%d %b %Y, %H:%M UTC"),
    )
    file_name = f"{_slugify(data.video_metadata.title)}-{uuid.uuid4().hex[:8]}.pdf"
    file_path = os.path.join(OUTPUT_DIR, file_name)
    try:
        HTML(string=html_content).write_pdf(file_path)
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"PDF generation failed: {exc}")
    return file_path


# ---------------------------------------------------------------------------
# 3. DOCX Generation — same data, brand-consistent styling
# ---------------------------------------------------------------------------
def _set_heading_style(paragraph, size=14, color=BRAND_COLOR_RGB):
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run("")
    run.font.size = Pt(size)
    run.font.bold = True
    run.font.color.rgb = color
    run.font.name = "Calibri"


def _add_bullets(doc, items, empty_text="No items provided."):
    if not items:
        p = doc.add_paragraph(empty_text)
        p.runs[0].font.italic = True
        p.runs[0].font.color.rgb = RGBColor(0x9A, 0xA0, 0xA6)
        return
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        for run in p.runs:
            run.font.name = "Calibri"
            run.font.size = Pt(11)


def _shade_cell(cell, hex_color):
    shd = cell._tc.get_or_add_tcPr().makeelement(qn("w:shd"), {})
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def _render_docx(data: SynopsisInput) -> str:
    try:
        meta = data.video_metadata
        summary = data.summary

        doc = Document()
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.font.color.rgb = TEXT_COLOR_RGB

        # Brand strip
        header_p = doc.add_paragraph()
        hr = header_p.add_run("VIDEO SYNOPSIS")
        hr.font.size = Pt(10)
        hr.font.bold = True
        hr.font.color.rgb = BRAND_COLOR_RGB
        dr = header_p.add_run(f"   |   Generated {datetime.utcnow().strftime('%d %b %Y, %H:%M UTC')}")
        dr.font.size = Pt(9)
        dr.font.color.rgb = MUTED_RGB

        # Title
        title = doc.add_heading(meta.title, level=1)
        _set_heading_style(title, size=20, color=RGBColor(0x20, 0x21, 0x24))

        # Metadata table
        rows = [("Channel", meta.channel_name)] if meta.channel_name else []
        rows.append(("Source URL", meta.video_url))
        meta_table = doc.add_table(rows=len(rows), cols=2)
        meta_table.alignment = WD_TABLE_ALIGNMENT.LEFT
        for i, (label, value) in enumerate(rows):
            label_cell, value_cell = meta_table.rows[i].cells
            label_cell.width = Inches(1.3)
            value_cell.width = Inches(5.0)
            _shade_cell(label_cell, "F8F9FA")
            lr = label_cell.paragraphs[0].add_run(label)
            lr.font.bold = True
            lr.font.size = Pt(10)
            vr = value_cell.paragraphs[0].add_run(value)
            vr.font.size = Pt(10)
        doc.add_paragraph()

        # Overall synopsis
        h = doc.add_heading("Overall Synopsis", level=2)
        _set_heading_style(h)
        p = doc.add_paragraph(summary.basic_summary.overall_synopsis)
        p.runs[0].font.size = Pt(11)

        # Topics covered
        h = doc.add_heading(summary.topics_covered.title, level=2)
        _set_heading_style(h)
        _add_bullets(doc, summary.topics_covered.topics, "No topics provided.")

        # Key insights
        h = doc.add_heading("Key Insights", level=2)
        _set_heading_style(h)
        _add_bullets(doc, summary.detailed_summary.key_insights, "No insights provided.")

        # Topic breakdown
        if summary.detailed_summary.topic_breakdown:
            h = doc.add_heading("Topic Breakdown", level=2)
            _set_heading_style(h)
            for item in summary.detailed_summary.topic_breakdown:
                sub = doc.add_heading(item.topic, level=3)
                _set_heading_style(sub, size=11, color=BRAND_COLOR_RGB)
                ep = doc.add_paragraph(item.explanation)
                ep.runs[0].font.size = Pt(10.5)

        # Action items
        h = doc.add_heading("Action Items", level=2)
        _set_heading_style(h)
        _add_bullets(doc, summary.detailed_summary.action_items, "No action items provided.")

        # Closing note
        h = doc.add_heading("Closing Note", level=2)
        _set_heading_style(h)
        cp = doc.add_paragraph(summary.closing_note)
        cp.runs[0].font.italic = True

        file_name = f"{_slugify(meta.title)}-{uuid.uuid4().hex[:8]}.docx"
        file_path = os.path.join(OUTPUT_DIR, file_name)
        doc.save(file_path)
        return file_path

    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"DOCX generation failed: {exc}")


# ---------------------------------------------------------------------------
# 4. Routes
#
#    PRODUCTION PATH (what the frontend actually calls):
#      GET /api/v1/synopsis/{synopsis_id}/download?format=pdf|docx
#      -> pulls the document straight out of MongoDB and exports it.
#
#    TESTING / INTEGRATION PATH (for you, M7, and Postman):
#      POST /api/export/pdf, /api/export/docx
#      -> takes the exact JSON M5 produces, no DB needed.
# ---------------------------------------------------------------------------
@app.get("/api/v1/synopsis/{synopsis_id}/download")
async def download_synopsis(synopsis_id: str, format: str, background_tasks: BackgroundTasks):
    if format not in ("pdf", "docx"):
        raise HTTPException(status_code=400, detail="format must be 'pdf' or 'docx'.")

    data = await _fetch_synopsis(synopsis_id)

    if format == "pdf":
        file_path = _render_pdf(data)
        media_type = "application/pdf"
    else:
        file_path = _render_docx(data)
        media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    background_tasks.add_task(_cleanup_file, file_path)
    return FileResponse(
        file_path,
        media_type=media_type,
        filename=os.path.basename(file_path),
        background=background_tasks,
    )


@app.post("/api/export/pdf")
def export_pdf(data: SynopsisInput, background_tasks: BackgroundTasks):
    file_path = _render_pdf(data)
    background_tasks.add_task(_cleanup_file, file_path)
    return FileResponse(
        file_path,
        media_type="application/pdf",
        filename=os.path.basename(file_path),
        background=background_tasks,
    )


@app.post("/api/export/docx")
def export_docx(data: SynopsisInput, background_tasks: BackgroundTasks):
    file_path = _render_docx(data)
    background_tasks.add_task(_cleanup_file, file_path)
    return FileResponse(
        file_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=os.path.basename(file_path),
        background=background_tasks,
    )
