"""DOCX rendering export functionality and styling helpers."""
import os
import uuid
from datetime import datetime
from fastapi import HTTPException

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

from app.schemas import SynopsisInput
from app.config import BRAND_COLOR_RGB, TEXT_COLOR_RGB, MUTED_RGB, OUTPUT_DIR
from app.utils import _slugify

def _set_heading_style(paragraph, size=14, color=BRAND_COLOR_RGB):
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run("")
    run.font.size = Pt(size)
    run.font.bold = True
    run.font.color.rgb = color
    run.font.name = "Calibri"

def _add_bullets(doc, items, empty_text="No items provided."):
    if not items:
        p = doc.add_paragraph(empty_text)
        p.runs[0].font.italic = True
        p.runs[0].font.color.rgb = RGBColor(0x9A, 0xA0, 0xA6)
        return
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        for run in p.runs:
            run.font.name = "Calibri"
            run.font.size = Pt(11)

def _shade_cell(cell, hex_color):
    shd = cell._tc.get_or_add_tcPr().makeelement(qn("w:shd"), {})
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)

def _render_docx(data: SynopsisInput) -> str:
    try:
        meta = data.video_metadata
        summary = data.summary

        doc = Document()
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.font.color.rgb = TEXT_COLOR_RGB

        # Brand strip
        header_p = doc.add_paragraph()
        hr = header_p.add_run("VIDEO SYNOPSIS")
        hr.font.size = Pt(10)
        hr.font.bold = True
        hr.font.color.rgb = BRAND_COLOR_RGB
        dr = header_p.add_run(f"   |   Generated {datetime.utcnow().strftime('%d %b %Y, %H:%M UTC')}")
        dr.font.size = Pt(9)
        dr.font.color.rgb = MUTED_RGB

        # Title
        title = doc.add_heading(meta.title, level=1)
        _set_heading_style(title, size=20, color=RGBColor(0x20, 0x21, 0x24))

        # Metadata table
        rows = [("Channel", meta.channel_name)] if meta.channel_name else []
        rows.append(("Source URL", meta.video_url))
        meta_table = doc.add_table(rows=len(rows), cols=2)
        meta_table.alignment = WD_TABLE_ALIGNMENT.LEFT
        for i, (label, value) in enumerate(rows):
            label_cell, value_cell = meta_table.rows[i].cells
            label_cell.width = Inches(1.3)
            value_cell.width = Inches(5.0)
            _shade_cell(label_cell, "F8F9FA")
            lr = label_cell.paragraphs[0].add_run(label)
            lr.font.bold = True
            lr.font.size = Pt(10)
            vr = value_cell.paragraphs[0].add_run(value)
            vr.font.size = Pt(10)
        doc.add_paragraph()

        # Overall synopsis
        h = doc.add_heading("Overall Synopsis", level=2)
        _set_heading_style(h)
        p = doc.add_paragraph(summary.basic_summary.overall_synopsis)
        p.runs[0].font.size = Pt(11)

        # Topics covered
        h = doc.add_heading(summary.topics_covered.title, level=2)
        _set_heading_style(h)
        _add_bullets(doc, summary.topics_covered.topics, "No topics provided.")

        # Key insights
        h = doc.add_heading("Key Insights", level=2)
        _set_heading_style(h)
        _add_bullets(doc, summary.detailed_summary.key_insights, "No insights provided.")

        # Topic breakdown
        if summary.detailed_summary.topic_breakdown:
            h = doc.add_heading("Topic Breakdown", level=2)
            _set_heading_style(h)
            for item in summary.detailed_summary.topic_breakdown:
                sub = doc.add_heading(item.topic, level=3)
                _set_heading_style(sub, size=11, color=BRAND_COLOR_RGB)
                ep = doc.add_paragraph(item.explanation)
                ep.runs[0].font.size = Pt(10.5)

        # Action items
        h = doc.add_heading("Action Items", level=2)
        _set_heading_style(h)
        _add_bullets(doc, summary.detailed_summary.action_items, "No action items provided.")

        # Closing note
        h = doc.add_heading("Closing Note", level=2)
        _set_heading_style(h)
        cp = doc.add_paragraph(summary.closing_note)
        cp.runs[0].font.italic = True

        file_name = f"{_slugify(meta.title)}-{uuid.uuid4().hex[:8]}.docx"
        file_path = os.path.join(OUTPUT_DIR, file_name)
        doc.save(file_path)
        return file_path

    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"DOCX generation failed: {exc}")
